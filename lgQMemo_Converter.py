import os
import glob
import zipfile
import json
from docx import Document
import shutil
import re

nowdir = os.getcwd()


print("LG Q메모 파일(.lqm) 변환기_v1.0(2209)\n제작: 파다기(https://blog.naver.com/padagi5678/222872694263)\n\n")
print("워드 파일(.docx)로 변환은 항상 가능하며, 메모장 파일(.txt)로는 일부 특수문자가 있을 경우 변환 실패할 수 있습니다.(실패한 경우 알림)\n워드 파일(.docx)에서는 메모에 있던 이미지도 함께 변환해드리니 .txt 파일말고 .docx 파일 사용을 권장합니다.")
print("변환이 끝나면 .lqm 파일의 확장자는 .zip 파일로 바뀝니다.(.lqm)\n")
s = input("엔터를 누르면 시작..")
print("-"*20)

try:
    os.mkdir("docx")
    os.mkdir("txt")
except FileExistsError:
    pass

files = glob.glob("./*.lqm")
for name in files:
    if not os.path.isdir(name):
        src = os.path.splitext(name)
        os.rename(name,src[0]+'.zip') #zip로

        memo_zip = zipfile.ZipFile(src[0]+'.zip')
        memo_zip.extractall(src[0])
        memo_zip.close() #압축해제

        files2 = glob.glob(f"{src[0]}/memoinfo.jlqm")
        src2 = os.path.splitext(files2[0])
        os.rename(files2[0],src2[0]+'.json')

        memoname = src[0][2:]

        with open(nowdir+"\\"+memoname+"\\memoinfo.json","r",encoding="UTF8") as f:
            json_data = json.load(f)

        doc = Document()
        
        text = "" #txt파일 저장용
        
        text_list = json_data["MemoObjectList"]
        for t in text_list:
            try:
                text += t["DescRaw"]
                
                doc.add_paragraph(t["DescRaw"])
            except KeyError:
                try:
                    filedir = t["FileName"] #사진,영상 파일이 있는 부분
                    doc.add_picture(f"./{memoname}/images/{filedir}")
                except:
                    print("error")
                    pass

        text = re.sub(r"[\xa0|\u200b]", " ", text)
        f = open(f"{nowdir}\\txt\\{memoname}.txt","w")
        try:
            f.write(text)
            print(f"저장완료-{memoname}")
        except:
            print(f"! {memoname}.txt 저장실패. 워드파일은 정상 저장.")
        f.close()

        doc.save(f'./docx/{memoname}.docx')

        shutil.rmtree(f"./{memoname}")
        
print("="*20)
e = input("완료되었습니다. docx 폴더와 txt 폴더를 확인해주세요.\n엔터를 눌러서 창 종료..")
